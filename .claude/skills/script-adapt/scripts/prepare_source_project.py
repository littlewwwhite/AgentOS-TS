#!/usr/bin/env python3
# input: --source-path pointing to a source file (txt/md/docx/xlsx/pdf)
# output: JSON with projectName, projectPath, originalSourcePath, sourceTextPath
# pos: CLI entry point for preparing a source project workspace

"""
Prepare source material for the script-adapt pipeline.

Converts any supported file format into PROJECT_DIR/source.txt (markdown),
ready for Phase 1 consumption.

Supported formats:
  - .txt / .md     → copy as-is
  - .docx          → extract paragraphs + embedded Excel tables as markdown
  - .xlsx          → convert all sheets to markdown tables
  - .pdf           → extract text (requires pdfplumber)

Dependencies (install as needed):
  - python-docx    → for .docx
  - openpyxl       → for .xlsx / embedded Excel in .docx
  - pdfplumber     → for .pdf

Usage:
  python3 prepare_source_project.py --source-path <file> --workspace-path "${PROJECT_DIR}"
"""

import json
import sys
import re
import argparse
import zipfile
import tempfile
from pathlib import Path
from typing import List

REPO_ROOT = Path(__file__).resolve().parents[4]
sys.path.insert(0, str(REPO_ROOT / "scripts"))

from pipeline_state import ensure_state, update_stage


# ---------- Format converters ----------

def convert_txt(source: Path) -> str:
    return source.read_text(encoding='utf-8')


def convert_docx(source: Path) -> tuple:
    """Extract all content from docx: paragraphs + embedded Excel sheets.

    Returns (content_str, has_storyboard_bool).
    """
    import docx

    doc = docx.Document(str(source))
    parts: List[str] = []

    # 1. Paragraphs
    for p in doc.paragraphs:
        text = p.text.strip()
        if text:
            parts.append(text)

    # 2. Tables in docx body
    for table in doc.tables:
        md = _table_to_markdown(table)
        if md:
            parts.append(md)

    # 3. Embedded xlsx files
    embedded = _extract_embedded_xlsx(source)
    for label, md in embedded:
        parts.append(f"\n## {label}\n")
        parts.append(md)

    return '\n\n'.join(parts), len(embedded) > 0


def convert_xlsx(source: Path) -> str:
    """Convert all sheets in an xlsx to merged-line markdown."""
    import openpyxl
    wb = openpyxl.load_workbook(str(source))
    parts: List[str] = []

    for sname in wb.sheetnames:
        ws = wb[sname]
        rows = _sheet_to_rows(ws)
        if rows:
            parts.append(f"## {sname}\n")
            parts.append(_rows_to_merged_lines(rows))

    return '\n\n'.join(parts)


def convert_pdf(source: Path) -> str:
    """Extract text from PDF pages."""
    import pdfplumber
    parts: List[str] = []

    with pdfplumber.open(str(source)) as pdf:
        for i, page in enumerate(pdf.pages, 1):
            text = page.extract_text()
            if text and text.strip():
                parts.append(text.strip())

    return '\n\n'.join(parts)


# ---------- Helpers ----------

def _table_to_markdown(table) -> str:
    """Convert a python-docx Table to merged-line markdown."""
    rows: List[List[str]] = []
    for row in table.rows:
        cells = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
        if any(cells):
            rows.append(cells)
    return _rows_to_merged_lines(rows) if rows else ''


def _rows_to_merged_lines(rows: List[List[str]]) -> str:
    """Merge each row into a single flat line: scene and visual joined by comma.

    Example output:
        🎬 第 1 集 🎬
        1. 修车厂，生锈的双手拧紧螺丝 SFX: Metal wrench clinking
        2. 修车厂，Liam擦去额头汗水 Liam: Just another rotten day.
        3. 修车厂外，一辆越野车疯狂加速 SFX: Engine roaring
    """
    if not rows:
        return ''

    lines: List[str] = []
    prev_scene = ''

    for row in rows[1:]:
        non_empty = [c for c in row if c]
        if not non_empty:
            continue

        first = row[0] if row else ''

        # Episode header
        if '🎬' in first or re.search(r'第\s*\d+\s*集', first):
            lines.append(f"\n{first}")
            continue

        # Scene (col 1), inherit from previous if empty
        scene = row[1] if len(row) > 1 else ''
        if not scene:
            scene = prev_scene
        prev_scene = scene

        # Build: "shot. scene，visual dialogue ..."
        parts = []
        if first:
            parts.append(f"{first}.")

        # Merge scene + remaining columns with comma
        remaining = [cell for cell in row[2:] if cell]
        if scene and remaining:
            parts.append(f"{scene}，{'  '.join(remaining)}")
        elif scene:
            parts.append(scene)
        else:
            parts.extend(remaining)

        lines.append(' '.join(parts))

    return '\n'.join(lines)


def _sheet_to_rows(ws) -> List[List[str]]:
    """Extract non-empty rows from an openpyxl worksheet."""
    rows: List[List[str]] = []
    for row in ws.iter_rows(values_only=True):
        cells = [str(c).strip() if c is not None else '' for c in row]
        if any(cells):
            rows.append(cells)
    return rows


def _extract_embedded_xlsx(docx_path: Path) -> List[tuple]:
    """Extract embedded xlsx from docx, return [(label, markdown_table)]."""
    results: List[tuple] = []

    try:
        with zipfile.ZipFile(docx_path) as z:
            xlsx_entries = sorted(
                n for n in z.namelist()
                if 'embeddings' in n and n.endswith('.xlsx')
            )

            if not xlsx_entries:
                return results

            import openpyxl

            with tempfile.TemporaryDirectory() as tmp:
                for idx, entry in enumerate(xlsx_entries, 1):
                    z.extract(entry, tmp)
                    xlsx_path = Path(tmp) / entry
                    wb = openpyxl.load_workbook(str(xlsx_path))

                    for sname in wb.sheetnames:
                        ws = wb[sname]
                        rows = _sheet_to_rows(ws)
                        if rows:
                            label = f"Embedded Sheet {idx}"
                            results.append((label, _rows_to_merged_lines(rows)))
    except zipfile.BadZipFile:
        pass

    return results


# ---------- Main ----------

CONVERTERS = {
    '.txt': convert_txt,
    '.md': convert_txt,
    '.docx': convert_docx,
    '.xlsx': convert_xlsx,
    '.pdf': convert_pdf,
}


def prepare_source_project(source_path: str, workspace_path: str) -> dict:
    resolved = Path(source_path).resolve()
    if not resolved.is_file():
        raise ValueError(f"Source path is not a file: {resolved}")

    ext = resolved.suffix.lower()
    converter = CONVERTERS.get(ext)
    if not converter:
        raise ValueError(
            f"Unsupported file format: {ext}. "
            f"Supported: {', '.join(CONVERTERS.keys())}"
        )

    project_name = resolved.stem
    project_path = Path(workspace_path).resolve()
    source_text_path = project_path / "source.txt"

    project_path.mkdir(parents=True, exist_ok=True)

    # Convert and write
    result = converter(resolved)
    # convert_docx returns (content, has_storyboard), others return str
    if isinstance(result, tuple):
        content, has_storyboard = result
    else:
        content = result
        has_storyboard = False

    source_text_path.write_text(content, encoding='utf-8')
    ensure_state(str(project_path))
    update_stage(
        str(project_path),
        "SCRIPT",
        "running",
        next_action="review SCRIPT",
    )

    # mode: "storyboard" = format conversion, "novel" = creative adaptation
    mode = 'storyboard' if has_storyboard else 'novel'

    return {
        "projectName": project_name,
        "projectPath": str(project_path),
        "originalSourcePath": str(resolved),
        "sourceTextPath": str(source_text_path),
        "format": ext,
        "mode": mode,
        "chars": len(content),
    }


if __name__ == "__main__":
    parser = argparse.ArgumentParser(
        description="Prepare source material for script-adapt pipeline"
    )
    parser.add_argument(
        "--source-path", required=True,
        help="Path to the source file (txt/md/docx/xlsx/pdf)"
    )
    parser.add_argument(
        "--workspace-path", default="workspace",
        help="Path to workspace directory (default: workspace)"
    )
    args = parser.parse_args()

    try:
        result = prepare_source_project(args.source_path, args.workspace_path)
        print(json.dumps(result, ensure_ascii=False, indent=2))
    except Exception as e:
        print(json.dumps({"error": str(e)}), file=sys.stderr)
        sys.exit(1)
